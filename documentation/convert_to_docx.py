import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def apply_text_formatting(paragraph, text):
    # Parse Markdown inline elements: **bold**, _italic_, *italic*, `code`
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner_text = part[2:-2]
            sub_parts = re.split(r'(\*.*?\*|_.*?_)', inner_text)
            for sub_part in sub_parts:
                if (sub_part.startswith('*') and sub_part.endswith('*')) or (sub_part.startswith('_') and sub_part.endswith('_')):
                    run = paragraph.add_run(sub_part[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(sub_part)
                    run.bold = True
        else:
            sub_parts = re.split(r'(\*.*?\*|_.*?_)', part)
            for sub_part in sub_parts:
                if (sub_part.startswith('*') and sub_part.endswith('*')) or (sub_part.startswith('_') and sub_part.endswith('_')):
                    sub_sub = re.split(r'(\`.*?\`)', sub_part[1:-1])
                    for ss in sub_sub:
                        if ss.startswith('`') and ss.endswith('`'):
                            run = paragraph.add_run(ss[1:-1])
                            run.italic = True
                            run.font.name = 'Times New Roman'
                        else:
                            run = paragraph.add_run(ss)
                            run.italic = True
                else:
                    sub_sub = re.split(r'(\`.*?\`)', sub_part)
                    for ss in sub_sub:
                        if ss.startswith('`') and ss.endswith('`'):
                            run = paragraph.add_run(ss[1:-1])
                            run.font.name = 'Times New Roman'
                        else:
                            paragraph.add_run(ss)

def process_table(doc, table_lines):
    if len(table_lines) < 2:
        return
    header_line = table_lines[0]
    data_lines = table_lines[2:] if len(table_lines) > 2 else []
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
    cols = len(headers)
    
    table = doc.add_table(rows=len(data_lines) + 1, cols=cols)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for col_idx, text in enumerate(headers):
        hdr_cells[col_idx].text = ''
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        set_cell_background(hdr_cells[col_idx], "F2F2F2")
        
    for row_idx, row_text in enumerate(data_lines):
        cells_text = [cell.strip() for cell in row_text.split('|')[1:-1]]
        while len(cells_text) < cols:
            cells_text.append('')
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(cells_text[:cols]):
            row_cells[col_idx].text = ''
            p = row_cells[col_idx].paragraphs[0]
            apply_text_formatting(p, text)
            
    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

def write_markdown_to_doc(doc, md_path, is_first):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if not stripped:
            if in_table:
                process_table(doc, table_lines)
                table_lines = []
                in_table = False
            i += 1
            continue
            
        if stripped.startswith('|'):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            process_table(doc, table_lines)
            table_lines = []
            in_table = False
            
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            if not is_first:
                p.paragraph_format.page_break_before = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        elif stripped.startswith('## '):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif stripped.startswith('### '):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif stripped.startswith('#### '):
            text = stripped[5:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.bold = True
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:-2].strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        elif stripped.startswith('$$'):
            math_content = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('$$'):
                math_content.append(lines[i].strip())
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("\n".join(math_content))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
        elif stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('• '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run('•\t')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            apply_text_formatting(p, stripped[2:])
            for r in p.runs[1:]:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
        elif re.match(r'^\d+\.\s+', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            match = re.match(r'^(\d+)\.\s+', stripped)
            num = match.group(1)
            run = p.add_run(f'{num}.\t')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            apply_text_formatting(p, stripped[match.end():])
            for r in p.runs[1:]:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            apply_text_formatting(p, stripped)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                
        i += 1
        
    if in_table and table_lines:
        process_table(doc, table_lines)

# Compile Chapter 1-3
doc = Document()

# Set standard margins (1.25 left for binding, 1 inch for top, bottom, right)
for s in doc.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.25)
    s.right_margin = Inches(1)

# Write chapters sequentially into a single file
write_markdown_to_doc(doc, "drafts/Chapter 1.md", True)
write_markdown_to_doc(doc, "drafts/Chapter 2.md", False)
write_markdown_to_doc(doc, "drafts/Chapter 3.md", False)

try:
    doc.save("docx/Chapter 1-3 ver3.docx")
    print("Successfully generated docx/Chapter 1-3 ver3.docx")
except PermissionError:
    alt_name = "docx/Chapter 1-3 ver3 (updated).docx"
    doc.save(alt_name)
    print(f"Permission denied for docx/Chapter 1-3 ver3.docx. Saved as {alt_name} instead.")

# Compile References separately
doc_ref = Document()
for s in doc_ref.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.25)
    s.right_margin = Inches(1)

write_markdown_to_doc(doc_ref, "drafts/References.md", True)

try:
    doc_ref.save("docx/References.docx")
    print("Successfully generated docx/References.docx")
except PermissionError:
    alt_ref_name = "docx/References (updated).docx"
    doc_ref.save(alt_ref_name)
    print(f"Permission denied for docx/References.docx. Saved as {alt_ref_name} instead.")

# Compile Chapter 4 separately
doc_ch4 = Document()
for s in doc_ch4.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.25)
    s.right_margin = Inches(1)

write_markdown_to_doc(doc_ch4, "drafts/Chapter 4.md", True)

try:
    doc_ch4.save("docx/Chapter 4.docx")
    print("Successfully generated docx/Chapter 4.docx")
except PermissionError:
    alt_ch4_name = "docx/Chapter 4 (updated).docx"
    doc_ch4.save(alt_ch4_name)
    print(f"Permission denied for docx/Chapter 4.docx. Saved as {alt_ch4_name} instead.")
